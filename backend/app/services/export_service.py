from typing import Dict
import json
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY
from datetime import datetime
import os
import logging

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting OCR results to various formats"""

    @staticmethod
    def export_to_txt(document: Dict, output_path: str) -> str:
        """Export OCR results to plain text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"OCR Results - {document['original_filename']}\n")
                f.write(f"{'='*60}\n\n")
                f.write(f"Processed: {document['processed_at']}\n")
                f.write(f"Confidence: {document['confidence']:.2%}\n")
                f.write(f"Total Lines: {document['line_count']}\n")
                f.write(f"\n{'='*60}\n")
                f.write(f"EXTRACTED TEXT\n")
                f.write(f"{'='*60}\n\n")
                f.write(document['extracted_text'])

            logger.info(f"Exported to TXT: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"TXT export failed: {str(e)}")
            raise

    @staticmethod
    def export_to_json(document: Dict, output_path: str) -> str:
        """Export OCR results to JSON file with full metadata"""
        try:
            export_data = {
                "document_info": {
                    "id": document['id'],
                    "original_filename": document['original_filename'],
                    "file_size": document['file_size'],
                    "file_type": document['file_type'],
                    "processed_at": document['processed_at'],
                    "created_at": document['created_at']
                },
                "ocr_results": {
                    "extracted_text": document['extracted_text'],
                    "confidence": document['confidence'],
                    "line_count": document['line_count'],
                    "status": document['status']
                },
                "detailed_lines": document.get('ocr_lines', [])
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

            logger.info(f"Exported to JSON: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"JSON export failed: {str(e)}")
            raise

    @staticmethod
    def export_to_docx(document: Dict, output_path: str) -> str:
        """Export OCR results to Microsoft Word document"""
        try:
            doc = DocxDocument()

            # Add title
            title = doc.add_heading('OCR Results', 0)
            title.alignment = 1  # Center

            # Add metadata
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_data = [
                ('Original Filename', document['original_filename']),
                ('Processed Date', document['processed_at']),
                ('Confidence', f"{document['confidence']:.2%}"),
                ('Total Lines', str(document['line_count'])),
                ('Status', document['status'].upper())
            ]

            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = value

            doc.add_paragraph()

            # Add extracted text
            doc.add_heading('Extracted Text', level=1)
            text_para = doc.add_paragraph(document['extracted_text'])
            text_para.paragraph_format.line_spacing = 1.15

            # Add page breaks between sections if multiple pages
            if document.get('ocr_lines'):
                doc.add_page_break()
                doc.add_heading('Detailed Line Data', level=1)

                for idx, line in enumerate(document['ocr_lines'][:50], 1):  # Limit to 50 lines
                    line_text = f"{idx}. {line['text']} (Confidence: {line['confidence']:.2%})"
                    doc.add_paragraph(line_text, style='List Number')

            doc.save(output_path)
            logger.info(f"Exported to DOCX: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"DOCX export failed: {str(e)}")
            raise

    @staticmethod
    def export_to_pdf(document: Dict, output_path: str) -> str:
        """Export OCR results to PDF document"""
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)

            # Container for PDF elements
            story = []

            # Styles
            styles = getSampleStyleSheet()
            title_style = styles['Heading1']
            title_style.alignment = 1  # Center

            normal_style = styles['Normal']
            normal_style.fontSize = 11
            normal_style.leading = 14

            # Add title
            story.append(Paragraph("OCR Results", title_style))
            story.append(Spacer(1, 0.2*inch))

            # Add metadata
            metadata = [
                f"<b>Original Filename:</b> {document['original_filename']}",
                f"<b>Processed Date:</b> {document['processed_at']}",
                f"<b>Confidence:</b> {document['confidence']:.2%}",
                f"<b>Total Lines:</b> {document['line_count']}",
                f"<b>Status:</b> {document['status'].upper()}"
            ]

            for meta in metadata:
                story.append(Paragraph(meta, normal_style))

            story.append(Spacer(1, 0.3*inch))

            # Add extracted text
            story.append(Paragraph("<b>Extracted Text</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))

            # Split text into paragraphs and add to PDF
            text_paragraphs = document['extracted_text'].split('\n')
            for para in text_paragraphs:
                if para.strip():
                    story.append(Paragraph(para, normal_style))
                    story.append(Spacer(1, 0.1*inch))

            # Build PDF
            doc.build(story)
            logger.info(f"Exported to PDF: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"PDF export failed: {str(e)}")
            raise


# Global export service instance
export_service = ExportService()
